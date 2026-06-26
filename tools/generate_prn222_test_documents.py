from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "test-documents" / "PRN222"


CONTENT = {
    "english": {
        "folder": "english-only",
        "docx_name": "PRN222_Review_DotNetCore_EN.docx",
        "pdf_name": "PRN222_Lecture_RazorPages_EFCore_EN.pdf",
        "md_name": "PRN222_FAQ_RAG_Test_EN.md",
        "txt_name": "PRN222_Quiz_API_Middleware_EN.txt",
        "csv_name": "PRN222_TestCases_EN.csv",
        "docx_title": "PRN222 - .NET Core Review Material",
        "docx_subtitle": "Sample documents for uploading into EduChatbot and testing RAG behavior.",
        "learning_goals": "1. Learning Objectives",
        "core_topics": "2. Core Knowledge",
        "middleware": "Middleware Pipeline",
        "ef_core": "Entity Framework Core",
        "test_table": "3. Content Test Table",
        "review_questions": "4. Review Questions",
        "goals": [
            "Understand how ASP.NET Core handles requests through the middleware pipeline.",
            "Explain the Razor Pages model and how data binding works in PageModel classes.",
            "Perform CRUD operations with Entity Framework Core, DbContext, and migrations.",
            "Apply dependency injection to separate Razor Pages, BLL, and DAL responsibilities.",
        ],
        "overview": (
            "ASP.NET Core is a cross-platform web framework from Microsoft. In PRN222, students "
            "usually work with Razor Pages, Entity Framework Core, validation, authentication, "
            "and layered application design with BLL and DAL."
        ),
        "middleware_items": [
            "UseRouting selects the endpoint that matches the URL.",
            "UseAuthentication reads the current user's identity.",
            "UseAuthorization checks access rights for a page or handler.",
            "MapRazorPages maps incoming requests to Razor Pages in the Pages folder.",
        ],
        "ef_text": (
            "DbContext represents a working session with the database. DbSet<T> represents an "
            "entity collection. Migrations record schema changes and allow the database to be "
            "updated with the dotnet ef database update command."
        ),
        "table_headers": ["Topic", "Keyword", "Test Question", "Expected Answer"],
        "table_rows": [
            ["Razor Pages", "OnGet, OnPost, ModelState", "When is OnPost called?", "When a form sends a POST request to a Razor Page."],
            ["EF Core", "DbContext, DbSet, Migration", "What is a migration used for?", "Managing database schema changes over time."],
            ["Dependency Injection", "Scoped, Transient, Singleton", "Which lifetime should DbContext use?", "Scoped, because each request needs its own working session."],
            ["Validation", "DataAnnotations, ModelState", "How do you check valid form data?", "Use ModelState.IsValid inside the handler."],
        ],
        "questions": [
            "Compare Razor Pages and MVC Controllers.",
            "Explain AddScoped, AddTransient, and AddSingleton.",
            "Write the steps for creating the first database migration.",
            "What should a PageModel return when ModelState.IsValid is false?",
            "Why should the UI avoid directly using DbContext when the project has BLL and DAL layers?",
        ],
        "pdf_title": "PRN222 - Sample Lecture: Razor Pages and EF Core",
        "pdf_purpose": "Purpose: this PDF is used to test document upload, page splitting, text extraction, and chatbot Q&A.",
        "pdf_h1_request": "1. ASP.NET Core Request Lifecycle",
        "pdf_request_text": (
            "When a user sends a request to an ASP.NET Core application, the request passes through "
            "a middleware chain. Middleware order directly affects the result, especially with "
            "routing, authentication, and authorization."
        ),
        "pdf_bullets": [
            "UseStaticFiles serves static files such as CSS, JavaScript, and images.",
            "UseRouting selects an endpoint based on the URL and HTTP method.",
            "UseAuthentication creates a ClaimsPrincipal when the user is signed in.",
            "UseAuthorization checks policies, roles, or access permissions.",
            "MapRazorPages connects endpoints to PageModel classes.",
        ],
        "pdf_h1_razor": "2. Razor Pages",
        "pdf_razor_text": (
            "Each Razor Page usually includes a .cshtml file and a PageModel .cshtml.cs file. "
            "The OnGet handler processes GET requests, while OnPost processes POST forms. "
            "Form data can be bound to properties by using BindProperty."
        ),
        "pdf_h1_ef": "3. Entity Framework Core",
        "pdf_ef_text": (
            "EF Core helps developers work with a database through C# objects. Students should "
            "understand DbContext, DbSet, LINQ, migrations, and how to separate repository or "
            "service logic in layered applications."
        ),
        "pdf_table": [
            ["Command", "Meaning", "When to use"],
            ["dotnet ef migrations add InitialCreate", "Creates the first migration", "After defining entities and DbContext"],
            ["dotnet ef database update", "Updates the database schema", "After creating or changing a migration"],
            ["AddDbContext()", "Registers DbContext", "In Program.cs so DI can provide it per request"],
        ],
        "pdf_h1_chatbot": "4. Suggested Chatbot Questions",
        "pdf_chatbot_notes": [
            "If a learner asks 'What is OnPost used for?', the chatbot should answer that OnPost handles POST requests, usually from form submissions.",
            "If a learner asks 'What is the role of DbContext?', the chatbot should explain that DbContext manages a database working session and tracks entity changes.",
        ],
        "md": """# PRN222 FAQ for RAG

## Razor Pages

**Question:** What is a PageModel used for?

**Answer:** A PageModel contains handlers, binding properties, and processing logic for a Razor Page.

## EF Core

**Question:** What does DbSet<T> represent?

**Answer:** DbSet<T> represents a collection of entities in a DbContext and usually maps to a database table.

## Middleware

**Question:** Why is middleware order important?

**Answer:** Because requests pass through middleware in registration order. If authorization runs before authentication, the application may not identify the user correctly.
""",
        "txt": """PRN222 - Short Quiz

1. What is ASP.NET Core middleware?
Answer: A component that processes requests and responses in the pipeline.

2. How are OnGet and OnPost different?
Answer: OnGet handles GET requests; OnPost handles POST requests, usually from form submissions.

3. What is ModelState.IsValid used for?
Answer: It checks whether binding and validation attributes produced valid data.

4. When do you need to create a migration?
Answer: When the entity schema changes and must be synchronized with the database.
""",
        "csv": """id,topic,input_question,expected_keyword
1,Razor Pages,What is OnGet?,GET handler
2,Razor Pages,When should OnPost be used?,POST form
3,EF Core,What is a migration used for?,database schema
4,DI,How is AddScoped different from Singleton?,lifetime
5,Validation,What does ModelState.IsValid mean?,validation
""",
    },
    "vietnamese": {
        "folder": "vietnamese-only",
        "docx_name": "PRN222_OnTap_DotNetCore_VI.docx",
        "pdf_name": "PRN222_BaiGiang_RazorPages_EFCore_VI.pdf",
        "md_name": "PRN222_FAQ_RAG_Test_VI.md",
        "txt_name": "PRN222_Quiz_API_Middleware_VI.txt",
        "csv_name": "PRN222_TestCases_VI.csv",
        "docx_title": "PRN222 - Tài liệu ôn tập .NET Core",
        "docx_subtitle": "Bộ tài liệu mẫu dùng để upload vào EduChatbot và kiểm thử hành vi RAG.",
        "learning_goals": "1. Mục tiêu học tập",
        "core_topics": "2. Kiến thức trọng tâm",
        "middleware": "Luồng middleware",
        "ef_core": "Entity Framework Core",
        "test_table": "3. Bảng kiểm thử nội dung",
        "review_questions": "4. Câu hỏi ôn tập",
        "goals": [
            "Hiểu cách ASP.NET Core xử lý request thông qua luồng middleware.",
            "Giải thích mô hình Razor Pages và cách binding dữ liệu trong PageModel.",
            "Thực hiện CRUD với Entity Framework Core, DbContext và migration.",
            "Áp dụng dependency injection để tách trách nhiệm giữa Razor Pages, BLL và DAL.",
        ],
        "overview": (
            "ASP.NET Core là framework web đa nền tảng của Microsoft. Trong môn PRN222, sinh viên "
            "thường làm việc với Razor Pages, Entity Framework Core, validation, authentication "
            "và thiết kế ứng dụng nhiều lớp theo hướng BLL và DAL."
        ),
        "middleware_items": [
            "UseRouting xác định endpoint phù hợp với URL.",
            "UseAuthentication đọc thông tin định danh của người dùng hiện tại.",
            "UseAuthorization kiểm tra quyền truy cập vào page hoặc handler.",
            "MapRazorPages ánh xạ request đến các Razor Page trong thư mục Pages.",
        ],
        "ef_text": (
            "DbContext đại diện cho một phiên làm việc với database. DbSet<T> đại diện cho một "
            "tập entity. Migration ghi lại thay đổi schema và cho phép cập nhật database bằng "
            "lệnh dotnet ef database update."
        ),
        "table_headers": ["Chủ đề", "Từ khóa", "Câu hỏi kiểm thử", "Đáp án kỳ vọng"],
        "table_rows": [
            ["Razor Pages", "OnGet, OnPost, ModelState", "Khi nào OnPost được gọi?", "Khi form gửi request POST đến Razor Page."],
            ["EF Core", "DbContext, DbSet, Migration", "Migration dùng để làm gì?", "Quản lý thay đổi schema database theo thời gian."],
            ["Dependency Injection", "Scoped, Transient, Singleton", "Nên đăng ký DbContext với lifetime nào?", "Scoped, vì mỗi request cần một phiên làm việc riêng."],
            ["Validation", "DataAnnotations, ModelState", "Làm sao kiểm tra dữ liệu form hợp lệ?", "Dùng ModelState.IsValid trong handler."],
        ],
        "questions": [
            "Phân biệt Razor Pages và MVC Controller.",
            "Giải thích ý nghĩa của AddScoped, AddTransient và AddSingleton.",
            "Viết các bước tạo migration đầu tiên cho database.",
            "Nếu ModelState.IsValid bằng false thì PageModel nên trả về kết quả nào?",
            "Tại sao không nên truy cập trực tiếp DbContext từ giao diện nếu dự án có BLL và DAL?",
        ],
        "pdf_title": "PRN222 - Bài giảng mẫu: Razor Pages và EF Core",
        "pdf_purpose": "Mục đích: file PDF này dùng để kiểm thử upload tài liệu, tách trang, trích xuất văn bản và hỏi đáp trên chatbot.",
        "pdf_h1_request": "1. Vòng đời request trong ASP.NET Core",
        "pdf_request_text": (
            "Khi người dùng gửi request đến ứng dụng ASP.NET Core, request sẽ đi qua chuỗi middleware. "
            "Thứ tự middleware ảnh hưởng trực tiếp đến kết quả xử lý, đặc biệt với routing, "
            "authentication và authorization."
        ),
        "pdf_bullets": [
            "UseStaticFiles phục vụ file tĩnh như CSS, JavaScript và hình ảnh.",
            "UseRouting chọn endpoint dựa trên URL và HTTP method.",
            "UseAuthentication tạo ClaimsPrincipal nếu người dùng đã đăng nhập.",
            "UseAuthorization kiểm tra policy, role hoặc quyền truy cập.",
            "MapRazorPages kết nối endpoint với các PageModel.",
        ],
        "pdf_h1_razor": "2. Razor Pages",
        "pdf_razor_text": (
            "Mỗi Razor Page thường gồm file .cshtml và file PageModel .cshtml.cs. Handler OnGet "
            "xử lý request GET, còn OnPost xử lý form POST. Dữ liệu form có thể bind vào property "
            "bằng BindProperty."
        ),
        "pdf_h1_ef": "3. Entity Framework Core",
        "pdf_ef_text": (
            "EF Core giúp lập trình viên làm việc với database bằng đối tượng C#. Sinh viên cần "
            "nắm DbContext, DbSet, LINQ, migration và cách tách repository hoặc service trong "
            "ứng dụng nhiều lớp."
        ),
        "pdf_table": [
            ["Lệnh", "Ý nghĩa", "Tình huống dùng"],
            ["dotnet ef migrations add InitialCreate", "Tạo migration đầu tiên", "Sau khi khai báo entity và DbContext"],
            ["dotnet ef database update", "Cập nhật schema database", "Sau khi tạo hoặc sửa migration"],
            ["AddDbContext()", "Đăng ký DbContext", "Trong Program.cs để DI cấp phát theo request"],
        ],
        "pdf_h1_chatbot": "4. Câu hỏi gợi ý cho chatbot",
        "pdf_chatbot_notes": [
            "Nếu người học hỏi 'OnPost dùng để làm gì?', chatbot nên trả lời rằng OnPost xử lý request POST, thường là submit form.",
            "Nếu người học hỏi 'DbContext có vai trò gì?', chatbot nên giải thích DbContext quản lý phiên làm việc với database và theo dõi thay đổi entity.",
        ],
        "md": """# PRN222 FAQ cho RAG

## Razor Pages

**Câu hỏi:** PageModel dùng để làm gì?

**Trả lời:** PageModel chứa handler, property binding và logic xử lý cho Razor Page.

## EF Core

**Câu hỏi:** DbSet<T> đại diện cho gì?

**Trả lời:** DbSet<T> đại diện cho một tập entity trong DbContext, thường ánh xạ đến một bảng trong database.

## Middleware

**Câu hỏi:** Vì sao thứ tự middleware quan trọng?

**Trả lời:** Vì request đi qua middleware theo thứ tự đăng ký. Nếu Authorization chạy trước Authentication, ứng dụng có thể không nhận diện được người dùng.
""",
        "txt": """PRN222 - Quiz ngắn

1. ASP.NET Core middleware là gì?
Đáp án: Thành phần xử lý request và response trong pipeline.

2. OnGet và OnPost khác nhau như thế nào?
Đáp án: OnGet xử lý request GET; OnPost xử lý request POST, thường từ form submit.

3. ModelState.IsValid dùng để làm gì?
Đáp án: Kiểm tra dữ liệu binding và validation attribute có hợp lệ hay không.

4. Khi nào cần tạo migration?
Đáp án: Khi schema entity thay đổi và cần đồng bộ xuống database.
""",
        "csv": """id,chủ_đề,câu_hỏi_đầu_vào,từ_khóa_kỳ_vọng
1,Razor Pages,OnGet là gì?,GET handler
2,Razor Pages,Khi nào dùng OnPost?,POST form
3,EF Core,Migration dùng để làm gì?,schema database
4,DI,AddScoped khác Singleton thế nào?,lifetime
5,Validation,ModelState.IsValid có ý nghĩa gì?,validation
""",
    },
}


def ensure_output_dirs() -> None:
    for content in CONTENT.values():
        (OUT / content["folder"]).mkdir(parents=True, exist_ok=True)


def add_docx_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def build_docx(content: dict[str, object]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, RGBColor(0, 0, 0)),
        ("Heading 2", 16, RGBColor(0, 0, 0)),
        ("Heading 3", 14, RGBColor(67, 67, 67)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(str(content["docx_title"]))
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph(str(content["docx_subtitle"]))
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading(str(content["learning_goals"]), level=1)
    for item in content["goals"]:
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading(str(content["core_topics"]), level=1)
    add_docx_paragraph(doc, str(content["overview"]))

    doc.add_heading(str(content["middleware"]), level=2)
    for item in content["middleware_items"]:
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading(str(content["ef_core"]), level=2)
    add_docx_paragraph(doc, str(content["ef_text"]))

    doc.add_heading(str(content["test_table"]), level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, header in enumerate(content["table_headers"]):
        table.rows[0].cells[idx].text = str(header)

    for row in content["table_rows"]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    doc.add_heading(str(content["review_questions"]), level=1)
    for question in content["questions"]:
        doc.add_paragraph(str(question), style="List Number")

    path = OUT / str(content["folder"]) / str(content["docx_name"])
    doc.save(path)
    return path


def register_pdf_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/tahoma.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("TestFont", str(candidate)))
            return "TestFont"
    return "Helvetica"


def bullet_list(items, style):
    return ListFlowable(
        [ListItem(Paragraph(str(item), style), leftIndent=14) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
    )


def build_pdf(content: dict[str, object]) -> Path:
    font = register_pdf_font()
    path = OUT / str(content["folder"]) / str(content["pdf_name"])
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=str(content["pdf_title"]),
    )

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontName=font,
        fontSize=22,
        leading=26,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=10,
    )
    h1 = ParagraphStyle(
        "H1Custom",
        parent=styles["Heading1"],
        fontName=font,
        fontSize=15,
        leading=18,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=14,
        spaceAfter=8,
    )
    body = ParagraphStyle(
        "BodyCustom",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=15,
        spaceAfter=7,
    )
    table_body = ParagraphStyle("TableBodyCustom", parent=body, fontSize=9.2, leading=12, spaceAfter=0)
    table_header = ParagraphStyle(
        "TableHeaderCustom",
        parent=table_body,
        textColor=colors.HexColor("#0B2545"),
    )

    story = [
        Paragraph(str(content["pdf_title"]), title),
        Paragraph(str(content["pdf_purpose"]), body),
        Paragraph(str(content["pdf_h1_request"]), h1),
        Paragraph(str(content["pdf_request_text"]), body),
        bullet_list(content["pdf_bullets"], body),
        Paragraph(str(content["pdf_h1_razor"]), h1),
        Paragraph(str(content["pdf_razor_text"]), body),
        Paragraph(str(content["pdf_h1_ef"]), h1),
        Paragraph(str(content["pdf_ef_text"]), body),
    ]

    data = []
    for row_index, row in enumerate(content["pdf_table"]):
        style = table_header if row_index == 0 else table_body
        data.append([Paragraph(str(cell), style) for cell in row])

    table = Table(data, colWidths=[2.05 * inch, 1.8 * inch, 3.05 * inch])
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFC7D5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    story.extend([Spacer(1, 8), table, Paragraph(str(content["pdf_h1_chatbot"]), h1)])
    for note in content["pdf_chatbot_notes"]:
        story.append(Paragraph(str(note), body))

    doc.build(story)
    return path


def build_text_files(content: dict[str, object]) -> list[Path]:
    folder = OUT / str(content["folder"])
    md = folder / str(content["md_name"])
    txt = folder / str(content["txt_name"])
    csv = folder / str(content["csv_name"])

    md.write_text(str(content["md"]), encoding="utf-8")
    txt.write_text(str(content["txt"]), encoding="utf-8")
    csv.write_text(str(content["csv"]), encoding="utf-8")
    return [md, txt, csv]


def main() -> None:
    ensure_output_dirs()
    paths = []
    for content in CONTENT.values():
        paths.extend([build_docx(content), build_pdf(content), *build_text_files(content)])
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
